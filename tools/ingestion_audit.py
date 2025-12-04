"""Ingestion audit utilities for the RAG sandbox.

This module performs three primary checks with verbose logging:
1. Extract text from PDFs, DOCX, or plaintext files while capturing per-file
   statistics (page counts, character counts, and empty-page detection).
2. Chunk, embed, and index the documents into a temporary Chroma collection.
3. Run a self-retrieval evaluation that queries the vector store with the
   original chunks to confirm that ingestion and embedding alignment are
   functioning correctly.

All output is recorded to both stdout and an audit log so investigators can
review the ingestion trace. The script favors deterministic, offline-friendly
execution and avoids external API calls.
"""
from __future__ import annotations

import argparse
import json
import logging
import random
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

import chromadb
import docx
import pypdfium2
from chromadb.utils import embedding_functions


AUDIT_LOG = Path("data/logs/ingestion_audit.log")
DEFAULT_CHROMA_PATH = Path("data/chroma_audit")
DEFAULT_COLLECTION = "ingestion_audit"
DEFAULT_EMBED_MODEL = "all-MiniLM-L6-v2"


@dataclass
class ExtractedDocument:
    """Container for a parsed document prior to chunking."""

    path: Path
    text: str
    page_count: int
    empty_pages: int


@dataclass
class ChunkRecord:
    """A single chunk ready for indexing and evaluation."""

    chunk_id: str
    source: str
    text: str


def configure_logging(verbose: bool = True) -> None:
    """Configure DEBUG-level logging to stdout and to the audit log file."""

    AUDIT_LOG.parent.mkdir(parents=True, exist_ok=True)
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(asctime)s | %(levelname)s | %(message)s",
        handlers=[
            logging.StreamHandler(),
            logging.FileHandler(AUDIT_LOG, encoding="utf-8"),
        ],
    )


def iter_files(inputs: Sequence[Path]) -> Iterable[Path]:
    """Yield individual files from a mix of file and directory paths."""

    for path in inputs:
        if path.is_dir():
            yield from path.rglob("*")
        else:
            yield path


def extract_text(path: Path) -> Tuple[str, int, int]:
    """Return (text, page_count, empty_pages) for the supplied file."""

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        doc = pypdfium2.PdfDocument(path)
        pages = []
        empty_pages = 0
        for page in doc:
            text = page.get_textpage().get_text_range()
            if not text.strip():
                empty_pages += 1
            pages.append(text)
        return "\n".join(pages), len(doc), empty_pages

    if suffix == ".docx":
        paragraphs = [p.text for p in docx.Document(path).paragraphs]
        text = "\n".join(p for p in paragraphs if p.strip())
        return text, len(paragraphs), sum(1 for p in paragraphs if not p.strip())

    text = path.read_text(encoding="utf-8", errors="ignore")
    return text, 1, 0 if text.strip() else 1


def chunk_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    """Split text into overlapping character windows for RAG indexing."""

    if chunk_size <= overlap:
        raise ValueError("chunk_size must be larger than overlap")

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks or [""]


def load_documents(paths: Sequence[Path]) -> List[ExtractedDocument]:
    """Parse provided files into ExtractedDocument objects with logging."""

    documents: List[ExtractedDocument] = []
    for path in iter_files(paths):
        if not path.is_file():
            continue
        logging.info(f"Reading {path}")
        try:
            text, page_count, empty_pages = extract_text(path)
        except Exception:  # noqa: BLE001
            logging.exception(f"Failed to extract text from {path}")
            continue

        doc = ExtractedDocument(path=path, text=text, page_count=page_count, empty_pages=empty_pages)
        logging.debug(
            "Parsed %s | %d chars | %d pages (%d empty)",
            path.name,
            len(text),
            page_count,
            empty_pages,
        )
        documents.append(doc)
    return documents


def build_embedding_function(model_name: str) -> embedding_functions.SentenceTransformerEmbeddingFunction:
    """Create a sentence-transformer embedding function with verbose logs."""

    logging.info(f"Loading embedding model: {model_name}")
    return embedding_functions.SentenceTransformerEmbeddingFunction(model_name=model_name)


def index_documents(
    client: chromadb.PersistentClient,
    collection_name: str,
    documents: Sequence[ExtractedDocument],
    embedder: embedding_functions.SentenceTransformerEmbeddingFunction,
    chunk_size: int,
    overlap: int,
    reset_collection: bool,
) -> List[ChunkRecord]:
    """Chunk and index documents into Chroma; return stored chunk metadata."""

    logging.info("Preparing Chroma collection")
    collection = client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},
        embedding_function=embedder,
    )

    if reset_collection:
        logging.debug("Clearing existing records from collection %s", collection_name)
        collection.delete(where={})

    chunk_records: List[ChunkRecord] = []
    for doc in documents:
        chunks = chunk_text(doc.text, chunk_size=chunk_size, overlap=overlap)
        ids = [f"{doc.path.name}-{i}" for i in range(len(chunks))]
        metadatas = [
            {"source": doc.path.name, "page_count": doc.page_count, "empty_pages": doc.empty_pages}
            for _ in chunks
        ]
        logging.debug("Indexing %d chunks for %s", len(chunks), doc.path.name)
        collection.add(ids=ids, documents=chunks, metadatas=metadatas)

        for chunk_id, text in zip(ids, chunks, strict=True):
            chunk_records.append(ChunkRecord(chunk_id=chunk_id, source=doc.path.name, text=text))

    logging.info("Indexed %d chunks across %d documents", len(chunk_records), len(documents))
    return chunk_records


def evaluate_self_retrieval(
    collection, chunks: Sequence[ChunkRecord],
    sample_size: int,
    k: int,
) -> dict:
    """Run self-retrieval and return metrics plus per-query diagnostics."""

    if not chunks:
        return {"samples": [], "success_rate": 0.0}

    sample = random.sample(chunks, k=min(sample_size, len(chunks)))
    results = []
    hits = 0

    for record in sample:
        query = record.text[:8000]
        res = collection.query(query_texts=[query], n_results=k)
        returned_ids = res.get("ids", [[]])[0]
        try:
            rank = returned_ids.index(record.chunk_id) + 1
            hits += 1
        except ValueError:
            rank = None
        results.append(
            {
                "chunk_id": record.chunk_id,
                "source": record.source,
                "hit": rank is not None,
                "rank": rank,
                "returned_ids": returned_ids,
            }
        )
        logging.debug(
            "Self-retrieval for %s | hit=%s | rank=%s | returned=%s",
            record.chunk_id,
            rank is not None,
            rank,
            returned_ids,
        )

    success_rate = hits / len(sample)
    logging.info("Self-retrieval success rate: %.2f", success_rate)
    return {"samples": results, "success_rate": success_rate}


def generate_report(
    documents: Sequence[ExtractedDocument],
    eval_results: dict,
    chunk_records: Sequence[ChunkRecord],
    output_path: Path,
) -> None:
    """Persist a JSON audit report for offline review."""

    report = {
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "documents": [
            {
                "path": str(doc.path),
                "characters": len(doc.text),
                "page_count": doc.page_count,
                "empty_pages": doc.empty_pages,
                "chunk_count": len([c for c in chunk_records if c.source == doc.path.name]),
            }
            for doc in documents
        ],
        "retrieval": eval_results,
    }

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
    logging.info("Saved audit report to %s", output_path)


def parse_args() -> argparse.Namespace:
    """Parse CLI options for the ingestion audit."""

    parser = argparse.ArgumentParser(description="Audit PDF ingestion and retrieval accuracy.")
    parser.add_argument("paths", nargs="+", type=Path, help="Files or folders to ingest")
    parser.add_argument("--chunk-size", type=int, default=1000, help="Chunk size in characters")
    parser.add_argument("--overlap", type=int, default=200, help="Overlap between chunks in characters")
    parser.add_argument("--embedding-model", default=DEFAULT_EMBED_MODEL, help="SentenceTransformer model name or path")
    parser.add_argument(
        "--chroma-path", default=DEFAULT_CHROMA_PATH, type=Path, help="Path for the temporary Chroma store"
    )
    parser.add_argument("--collection", default=DEFAULT_COLLECTION, help="Chroma collection name")
    parser.add_argument("--no-reset", action="store_true", help="Do not clear the collection before indexing")
    parser.add_argument("--samples", type=int, default=25, help="Number of chunks to probe during self-retrieval")
    parser.add_argument("--top-k", type=int, default=4, help="How many neighbors to request during self-retrieval")
    parser.add_argument(
        "--report", type=Path, default=Path("data/reports/ingestion_audit_report.json"), help="Where to write the JSON report"
    )
    return parser.parse_args()


def main() -> None:
    """Entry point that orchestrates extraction, indexing, and evaluation."""

    args = parse_args()
    configure_logging(verbose=True)

    documents = load_documents(args.paths)
    if not documents:
        logging.error("No readable documents found. Exiting.")
        return

    embedder = build_embedding_function(args.embedding_model)
    client = chromadb.PersistentClient(path=str(args.chroma_path))
    chunk_records = index_documents(
        client=client,
        collection_name=args.collection,
        documents=documents,
        embedder=embedder,
        chunk_size=args.chunk_size,
        overlap=args.overlap,
        reset_collection=not args.no_reset,
    )

    collection = client.get_collection(name=args.collection, embedding_function=embedder)
    eval_results = evaluate_self_retrieval(collection, chunk_records, sample_size=args.samples, k=args.top_k)
    generate_report(documents, eval_results, chunk_records, args.report)


if __name__ == "__main__":
    main()
