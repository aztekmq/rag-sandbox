import os
import shutil
from pathlib import Path
from typing import List
import gradio as gr
import pypdfium2
import requests
import chromadb
import docx
import json
import logging
from datetime import datetime

# â€”â€”â€”â€”â€”â€”â€” ADVANCED LOGGER â€”â€”â€”â€”â€”â€”â€”
LOGFILE = "rag_portal.log"
logging.basicConfig(
    filename=LOGFILE,
    level=logging.DEBUG,
    format="%(asctime)s | %(levelname)s | %(message)s",
)

console = logging.StreamHandler()
console.setLevel(logging.DEBUG)
console.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
logging.getLogger("").addHandler(console)

# â€”â€”â€”â€”â€”â€”â€” GRADIO BOOLEAN SCHEMA FIX (keeps it working) â€”â€”â€”â€”â€”â€”â€”
try:
    from gradio_client import utils as grc_utils
    orig = grc_utils.get_type

    def safe_get_type(s):
        return "any" if isinstance(s, bool) else orig(s)

    grc_utils.get_type = safe_get_type
except Exception:
    pass

os.environ["GRADIO_ANALYTICS_ENABLED"] = "false"

# â€”â€”â€”â€”â€”â€”â€” CONFIG â€”â€”â€”â€”â€”â€”â€”
UPLOAD_DIR = Path("data/rag_portal/uploads")
CHROMA_DIR = Path("data/rag_portal/chroma")
COLLECTION = "ibm_mq"

# ğŸ”¹ Read from environment (fall back to previous defaults)
EMBED_MODEL = os.getenv("EMBED_MODEL", "nomic-embed-text:latest")
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
DEFAULT_CHAT_MODEL = os.getenv("OLLAMA_CHAT_MODEL", "llama3.1:latest")  # must match /api/tags

logging.info(f"Using OLLAMA_URL={OLLAMA_URL}")
logging.info(f"Using EMBED_MODEL={EMBED_MODEL}")
logging.info(f"Default chat model={DEFAULT_CHAT_MODEL}")

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
CHROMA_DIR.mkdir(parents=True, exist_ok=True)

client = chromadb.PersistentClient(path=str(CHROMA_DIR))
collection = client.get_or_create_collection(
    name=COLLECTION, metadata={"hnsw:space": "cosine"}
)


def embed(text: str) -> List[float]:
    """Call Ollama embeddings endpoint; fall back to zeros if anything fails."""
    try:
        r = requests.post(
            f"{OLLAMA_URL}/api/embeddings",
            json={"model": EMBED_MODEL, "prompt": text},
            timeout=60,
        )
        r.raise_for_status()
        return r.json()["embedding"]
    except Exception:
        # Fallback dummy vector â€“ keeps the pipeline from crashing
        logging.exception("Embedding call failed; returning dummy vector.")
        return [0.0] * 768


def chunk_text(text: str) -> List[str]:
    """Chunk text into overlapping word blocks for RAG."""
    words = text.split()
    return [" ".join(words[i : i + 700]) for i in range(0, len(words), 600)] or [""]


# â€”â€”â€”â€”â€”â€”â€” INGEST â€” FIXED SAMEFILEERROR â€”â€”â€”â€”â€”â€”â€”
def ingest(files):
    if not files:
        return "No files selected"

    log = []
    for file_obj in files:  # â† file_obj is a Gradio File object
        original_path = Path(file_obj.name)  # temporary path
        dest_path = UPLOAD_DIR / original_path.name

        # If file is already in the right place â†’ skip copy
        if original_path.resolve() == dest_path.resolve():
            src = original_path
        else:
            src = shutil.copy2(original_path, dest_path)  # â† safe copy

        # Extract text
        try:
            if dest_path.suffix.lower() == ".pdf":
                text = "\n".join(
                    p.get_textpage().get_text_range()
                    for p in pypdfium2.PdfDocument(dest_path)
                )
            elif dest_path.suffix.lower() == ".docx":
                text = "\n".join(
                    p.text
                    for p in docx.Document(dest_path).paragraphs
                    if p.text.strip()
                )
            else:
                text = dest_path.read_text(encoding="utf-8")
        except Exception as e:
            logging.exception(f"Error reading {dest_path}")
            log.append(f"{dest_path.name} â†’ read error: {e}")
            continue

        chunks = chunk_text(text)
        ids = [f"{dest_path.name}_{i}" for i in range(len(chunks))]
        embeddings = [embed(c) for c in chunks]

        collection.add(
            ids=ids,
            documents=chunks,
            embeddings=embeddings,
            metadatas=[{"source": dest_path.name} for _ in chunks],
        )
        log.append(f"{dest_path.name} â†’ {len(chunks)} chunks indexed")

    return "\n".join(log)


# â€”â€”â€”â€”â€”â€”â€” MODEL LIST â€”â€”â€”â€”â€”â€”â€”
def get_models():
    """
    Return only chat-capable models for the dropdown.
    We exclude anything with 'embed' in the name.
    """
    try:
        r = requests.get(f"{OLLAMA_URL}/api/tags", timeout=5)
        r.raise_for_status()
        all_models = [m["name"] for m in r.json().get("models", [])]
        chat_models = [m for m in all_models if "embed" not in m.lower()]
        return chat_models or [DEFAULT_CHAT_MODEL]
    except Exception:
        logging.exception("Failed to fetch models from Ollama; using default.")
        return [DEFAULT_CHAT_MODEL]


# â€”â€”â€”â€”â€”â€”â€” CHAT â€”â€”â€”â€”â€”â€”â€”
def chat(message, history, model):
    """
    Fully instrumented chat function with:
      - Payload logging
      - Retrieval context logging
      - Embedding vector sizes
      - Full Ollama request/response bodies
      - Error detail logging
    """

    logging.info("â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€")
    logging.info("CHAT REQUEST")
    logging.info(f"Incoming user message: {message!r}")
    logging.info(f"Selected model: {model}")

    # Safety: force chat-capable model
    if model and "embed" in model.lower():
        logging.warning(f"User selected embedding model ({model}); overriding.")
        model = DEFAULT_CHAT_MODEL

    model = model or DEFAULT_CHAT_MODEL
    logging.info(f"Chat model in use: {model}")
    logging.info(f"OLLAMA_URL in chat(): {OLLAMA_URL}")

    # 1) Retrieve chunks
    q_emb = embed(message)
    logging.debug(f"Query embedding length: {len(q_emb)}")

    try:
        res = collection.query(query_embeddings=[q_emb], n_results=6)
        docs = res["documents"][0]
        metas = res["metadatas"][0]
        context = "\n\n".join(
            f"[Source: {m.get('source', 'unknown')}] {d[:1200]}"
            for d, m in zip(docs, metas)
        )
        logging.info(f"Retrieved {len(docs)} context chunks.")
    except Exception as e:
        logging.error(f"Retrieval failed: {e}")
        docs, metas = [], []
        context = "No indexed documents were found. Answer from general MQ knowledge."

    logging.debug("RAG CONTEXT SENT TO MODEL:\n" + context)

    # 2) Build payload
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an IBM MQ expert. Use ONLY the provided context "
                    "when available. Cite sources in brackets like [Source: file]."
                ),
            },
            {
                "role": "user",
                "content": f"Context:\n{context}\n\nQuestion: {message}",
            },
        ],
        "stream": False,
    }

    logging.info("OLLAMA PAYLOAD:")
    logging.debug(json.dumps(payload, indent=2))

    # 3) Send request
    url = f"{OLLAMA_URL}/api/chat"
    logging.info(f"Sending POST â†’ {url}")

    try:
        r = requests.post(url, json=payload, timeout=600)
        logging.info(f"Ollama HTTP status: {r.status_code}")
        logging.debug(f"Ollama raw response text:\n{r.text}")

        r.raise_for_status()

        data = r.json()
        logging.debug("Ollama parsed JSON:\n" + json.dumps(data, indent=2))

        answer = data.get("message", {}).get("content", "").strip()
        if not answer:
            logging.warning("Empty assistant message received from Ollama.")
            answer = "Received an empty response from Ollama."

    except (requests.ConnectionError, requests.Timeout) as e:
        logging.error(f"Connection error: {e}", exc_info=True)
        answer = (
            f"Ollama connection error: {e}\n\n"
            f"Make sure your Ollama server is running at {OLLAMA_URL}."
        )

    except requests.HTTPError as e:
        text = e.response.text if e.response is not None else "(no body)"
        logging.error(
            f"HTTP {e.response.status_code if e.response else '???'}:\n{text}",
            exc_info=True,
        )

        answer = (
            f"Ollama HTTP error {e.response.status_code if e.response else '???'}:\n\n"
            f"{text}"
        )

    except Exception as e:
        logging.error(f"Unexpected error: {e}", exc_info=True)
        answer = f"Unexpected error while calling Ollama: {e}"

    logging.info("Assistant reply: " + repr(answer))
    logging.info("â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€")

    return answer


# â€”â€”â€”â€”â€”â€”â€” UI â€”â€”â€”â€”â€”â€”â€”
with gr.Blocks(title="IBM MQ RAG") as demo:
    gr.Markdown("# IBM MQ RAG Portal â€” Final Working Version")
    gr.Markdown("**Embedding:** nomic-embed-text:latestâ€ƒ**LLM:** Ollama")

    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### Upload & Index")
            files = gr.File(
                file_count="multiple",
                file_types=[".pdf", ".docx", ".txt", ".md"],
            )
            btn = gr.Button("Ingest & Index", variant="primary")
            log = gr.Textbox(label="Log", lines=10)
            btn.click(ingest, files, log)

        with gr.Column(scale=2):
            gr.Markdown("### Ask Questions")
            model = gr.Dropdown(
                choices=get_models(),
                value=DEFAULT_CHAT_MODEL,
                label="Model",
            )
            gr.ChatInterface(chat, additional_inputs=model)

    # Refresh model list on load (e.g., after pulling new models)
    demo.load(lambda: gr.update(choices=get_models()), outputs=model)

# â€”â€”â€”â€”â€”â€”â€” LAUNCH â€”â€”â€”â€”â€”â€”â€”
demo.queue().launch(
    server_name="0.0.0.0",
    server_port=7863,
    share=False,
    inbrowser=False,
)